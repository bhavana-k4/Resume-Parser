"""
modules/extractor.py
Extracts clean text from PDF, DOCX, or plain-text resume files.
Also wraps raw strings (pasted job descriptions) in the same contract.
"""
import os, re, logging
from dataclasses import dataclass, field
from typing import Optional

log = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    text:      str  = ""
    raw_text:  str  = ""
    metadata:  dict = field(default_factory=dict)
    success:   bool = False
    error:     Optional[str] = None


# ── helpers ──────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    out, blanks = [], 0
    for line in [l.rstrip() for l in text.split('\n')]:
        if not line.strip():
            blanks += 1
            if blanks <= 2:
                out.append(line)
        else:
            blanks = 0
            out.append(line)
    return '\n'.join(out).strip()


def _meta(name, fmt, raw, extra=None):
    d = {"filename": os.path.basename(name), "format": fmt,
         "char_count": len(raw), "word_count": len(raw.split())}
    if extra:
        d.update(extra)
    return d


# ── format-specific ───────────────────────────────────────────────────────────

def _pdf(path):
    try:
        import pdfplumber
    except ImportError:
        return ExtractionResult(success=False, error="pip install pdfplumber")
    try:
        pages = []
        with pdfplumber.open(path) as pdf:
            n = len(pdf.pages)
            for pg in pdf.pages:
                t = pg.extract_text(x_tolerance=3, y_tolerance=3)
                if not (t and t.strip()):
                    ws = pg.extract_words()
                    t  = ' '.join(w['text'] for w in ws) if ws else ''
                if t and t.strip():
                    pages.append(t)
        if not pages:
            return ExtractionResult(success=False,
                error="No text found. Scanned PDF? Add OCR (pytesseract).")
        raw = '\n\n'.join(pages)
        return ExtractionResult(text=_clean(raw), raw_text=raw,
                                metadata=_meta(path, "pdf", raw, {"page_count": n}),
                                success=True)
    except Exception as e:
        return ExtractionResult(success=False, error=str(e))


def _docx(path):
    try:
        from docx import Document
    except ImportError:
        return ExtractionResult(success=False, error="pip install python-docx")
    try:
        doc   = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append('  |  '.join(cells))
        if not parts:
            return ExtractionResult(success=False, error="Empty document.")
        raw = '\n'.join(parts)
        return ExtractionResult(text=_clean(raw), raw_text=raw,
                                metadata=_meta(path, "docx", raw,
                                    {"para": len(doc.paragraphs), "tables": len(doc.tables)}),
                                success=True)
    except Exception as e:
        return ExtractionResult(success=False, error=str(e))


def _txt(path):
    for enc in ('utf-8', 'latin-1', 'cp1252'):
        try:
            raw = open(path, encoding=enc).read()
            return ExtractionResult(text=_clean(raw), raw_text=raw,
                                    metadata=_meta(path, "txt", raw), success=True)
        except UnicodeDecodeError:
            continue
        except Exception as e:
            return ExtractionResult(success=False, error=str(e))
    return ExtractionResult(success=False, error="Cannot decode file. Save as UTF-8.")


# ── public API ────────────────────────────────────────────────────────────────

SUPPORTED = {'.pdf': _pdf, '.docx': _docx, '.txt': _txt}


def extract_text(filepath: str) -> ExtractionResult:
    """Extract text from a saved file (PDF / DOCX / TXT)."""
    if not os.path.isfile(filepath):
        return ExtractionResult(success=False, error=f"File not found: {filepath}")
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in SUPPORTED:
        return ExtractionResult(success=False,
            error=f"Unsupported format '{ext}'. Use: {list(SUPPORTED)}")
    return SUPPORTED[ext](filepath)


def extract_from_string(text: str, label: str = "input") -> ExtractionResult:
    """Wrap a raw pasted string in an ExtractionResult."""
    if not text or not text.strip():
        return ExtractionResult(success=False, error="Empty input.")
    return ExtractionResult(text=_clean(text), raw_text=text,
                            metadata=_meta(label, "string", text), success=True)